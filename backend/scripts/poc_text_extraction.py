"""텍스트 추출 PoC — pdf/docx 더미 이력서 생성 → 추출 → 원본 비교.

실행: python scripts/poc_text_extraction.py
결과: scripts/poc_results/ 에 샘플 파일 + 결과 JSON
"""

import json
import random
import sys
import time
from pathlib import Path

SCRIPT_DIR = Path(__file__).parent
RESULTS_DIR = SCRIPT_DIR / "poc_results"
RESULTS_DIR.mkdir(exist_ok=True)

MATERIALS_DIR = SCRIPT_DIR / "seed" / "materials"

# ── 재료 로드 ──────────────────────────────────────────────────────────

with open(MATERIALS_DIR / "names.json", encoding="utf-8") as f:
    names_data = json.load(f)
with open(MATERIALS_DIR / "schools.json", encoding="utf-8") as f:
    schools_data = json.load(f)
with open(MATERIALS_DIR / "skill-patterns.json", encoding="utf-8") as f:
    skill_data = json.load(f)


def random_resume():
    last = random.choice(names_data["last"])
    first = random.choice(names_data["first"])
    name = last + first
    school_info = random.choice(schools_data)
    school = school_info["school"]
    major = random.choice(school_info["majors"])
    track = random.choice(skill_data)
    skills = random.sample(track["skills"], min(4, len(track["skills"])))
    years = random.randint(1, 10)
    email = f"{name.lower()}@example.com"

    text = f"""이력서

이름: {name}
이메일: {email}
연락처: 010-{random.randint(1000,9999)}-{random.randint(1000,9999)}

학력
{school} {major} 졸업

경력
{track['track']} 개발자 {years}년차

보유 기술
{', '.join(skills)}

자기소개
안녕하세요, {name}입니다. {school} {major}를 졸업하고 {track['track']} 분야에서 {years}년간 근무하였습니다. 주요 기술 스택으로는 {', '.join(skills[:2])}를 활용하여 다양한 프로젝트를 수행하였으며, 팀 협업과 문제 해결 능력을 갖추고 있습니다.
"""
    return name, text.strip()


# ── PDF 생성 ──────────────────────────────────────────────────────────

def create_pdf(text: str, path: Path):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    font_path = "C:/Windows/Fonts/malgun.ttf"
    pdf.add_font("malgun", "", font_path, uni=True)
    pdf.set_font("malgun", size=11)

    for line in text.split("\n"):
        pdf.cell(0, 8, line, new_x="LMARGIN", new_y="NEXT")

    pdf.output(str(path))


# ── DOCX 생성 ─────────────────────────────────────────────────────────

def create_docx(text: str, path: Path):
    from docx import Document

    doc = Document()
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    doc.save(str(path))


# ── 텍스트 추출 ──────────────────────────────────────────────────────

def extract_pdf(path: Path) -> str:
    import pdfplumber

    with pdfplumber.open(str(path)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages)


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


# ── 비교 ─────────────────────────────────────────────────────────────

def compare(original: str, extracted: str) -> dict:
    orig_lines = [l.strip() for l in original.split("\n") if l.strip()]
    ext_lines = [l.strip() for l in extracted.split("\n") if l.strip()]

    matched = 0
    for orig_line in orig_lines:
        if any(orig_line in ext_line or ext_line in orig_line for ext_line in ext_lines):
            matched += 1

    coverage = round(matched / len(orig_lines) * 100, 1) if orig_lines else 0
    return {
        "original_lines": len(orig_lines),
        "extracted_lines": len(ext_lines),
        "matched_lines": matched,
        "coverage_pct": coverage,
    }


# ── 실행 ─────────────────────────────────────────────────────────────

def main():
    results = []
    random.seed(42)

    print("=" * 60)
    print("Arda 텍스트 추출 PoC")
    print("=" * 60)

    # PDF 10건
    print("\n[PDF 생성 + 추출]")
    for i in range(10):
        name, text = random_resume()
        pdf_path = RESULTS_DIR / f"resume_{i+1:02d}_{name}.pdf"

        try:
            create_pdf(text, pdf_path)
            extracted = extract_pdf(pdf_path)
            cmp = compare(text, extracted)
            status = "SUCCESS" if cmp["coverage_pct"] >= 80 else "PARTIAL"
            print(f"  {i+1:2d}. {name:6s} PDF  → {status} ({cmp['coverage_pct']}%)")
        except Exception as e:
            status = "FAIL"
            cmp = {"error": str(e)}
            print(f"  {i+1:2d}. {name:6s} PDF  → FAIL: {e}")

        results.append({
            "id": i + 1,
            "name": name,
            "format": "pdf",
            "file": pdf_path.name,
            "status": status,
            **cmp,
        })

    # DOCX 5건
    print("\n[DOCX 생성 + 추출]")
    for i in range(5):
        name, text = random_resume()
        docx_path = RESULTS_DIR / f"resume_{i+11:02d}_{name}.docx"

        try:
            create_docx(text, docx_path)
            extracted = extract_docx(docx_path)
            cmp = compare(text, extracted)
            status = "SUCCESS" if cmp["coverage_pct"] >= 80 else "PARTIAL"
            print(f"  {i+11:2d}. {name:6s} DOCX → {status} ({cmp['coverage_pct']}%)")
        except Exception as e:
            status = "FAIL"
            cmp = {"error": str(e)}
            print(f"  {i+11:2d}. {name:6s} DOCX → FAIL: {e}")

        results.append({
            "id": i + 11,
            "name": name,
            "format": "docx",
            "file": docx_path.name,
            "status": status,
            **cmp,
        })

    # HWP — 라이브러리 테스트만
    print("\n[HWP 라이브러리 확인]")
    try:
        import hwp5
        print("  python-hwp5 설치됨 — 실제 hwp 파일로 테스트 필요")
        hwp_lib = "python-hwp5"
    except ImportError:
        try:
            import olefile
            print("  olefile 설치됨 — hwp 파싱 가능하나 불안정")
            hwp_lib = "olefile"
        except ImportError:
            print("  hwp 라이브러리 없음 — 수동 폴백 경로 사용")
            hwp_lib = "none"

    results.append({
        "id": 16,
        "name": "-",
        "format": "hwp",
        "file": "-",
        "status": "NOT_TESTED",
        "note": f"라이브러리: {hwp_lib}. 실제 hwp 파일로 별도 테스트 필요",
    })

    # 결과 저장
    report_path = RESULTS_DIR / "poc_report.json"
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    # 요약
    pdf_ok = sum(1 for r in results if r["format"] == "pdf" and r["status"] == "SUCCESS")
    docx_ok = sum(1 for r in results if r["format"] == "docx" and r["status"] == "SUCCESS")

    print("\n" + "=" * 60)
    print("요약")
    print("=" * 60)
    print(f"  PDF:  {pdf_ok}/10 성공")
    print(f"  DOCX: {docx_ok}/5  성공")
    print(f"  HWP:  별도 테스트 필요 (수동 폴백 확정)")
    print(f"\n  결과: {report_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
